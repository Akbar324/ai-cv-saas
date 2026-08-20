"""Live synthetic end-to-end persistence smoke test."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from tempfile import TemporaryDirectory

import boto3  # type: ignore[import-untyped]
from docx import Document

from backend.models.aws_settings import load_aws_settings
from backend.models.order import Order
from backend.services.cv_workflow import process_and_persist_cv

TEST_ORDER_ID = "smoke-e2e-order-001"
TEST_CUSTOMER_ID = "smoke-customer-001"


def create_synthetic_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Alex Morgan")
    document.add_paragraph("Dubai, UAE")
    document.add_paragraph("Quality Engineer")
    document.add_paragraph("Example Engineering LLC")
    document.add_paragraph("2021 - Present")
    document.add_paragraph("Perform technical inspections of telecom infrastructure.")
    document.add_paragraph("Review installation work against engineering standards.")
    document.add_paragraph("Coordinate with field teams to resolve quality issues.")
    document.add_paragraph("Prepare inspection reports.")
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor of Engineering")
    document.add_paragraph("Skills")
    document.add_paragraph("Quality Inspection")
    document.add_paragraph("Telecommunications")
    document.add_paragraph("Python")
    document.add_paragraph("AWS")

    document.save(str(path))


def main() -> None:
    aws_settings = load_aws_settings()

    session = boto3.Session(
        profile_name=aws_settings.profile,
        region_name=aws_settings.region,
    )

    dynamodb = session.resource("dynamodb")
    s3 = session.client("s3")

    table = dynamodb.Table(aws_settings.orders_table_name)

    now = datetime.now(UTC)

    order = Order(
        order_id=TEST_ORDER_ID,
        customer_id=TEST_CUSTOMER_ID,
        target_job_title="Junior Cloud Engineer",
        target_industry="Cloud Computing",
        created_at=now,
        updated_at=now,
    )

    print("Creating synthetic order in DynamoDB...")

    table.put_item(
        Item=order.model_dump(mode="json", exclude={"documents"}),
        ConditionExpression="attribute_not_exists(order_id)",
    )

    try:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "synthetic-candidate.docx"
            create_synthetic_docx(source_path)

            print("Running configured CV workflow...")

            result = process_and_persist_cv(
                order=order,
                source_path=source_path,
                job_description=(
                    "Seeking a junior cloud engineer with AWS knowledge, "
                    "basic Python skills, troubleshooting ability, and "
                    "experience working in technical environments."
                ),
                additional_customer_information=(
                    "Candidate is transitioning toward cloud engineering. "
                    "Do not invent production cloud experience."
                ),
            )

            print()
            print("Workflow completed.")
            print("Order status:", result.order_status.value)
            print("Processing status:", result.processing_status.value)
            print("AI provider:", result.ai_provider)
            print("AI model:", result.ai_model)
            print("CV version:", result.current_cv_version)
            print("Source key:", result.documents.source_s3_key)
            print("CV key:", result.documents.current_cv_s3_key)

            stored = table.get_item(
                Key={"order_id": TEST_ORDER_ID},
            ).get("Item")

            print()
            print("DynamoDB record exists:", stored is not None)

            source_key = result.documents.source_s3_key
            cv_key = result.documents.current_cv_s3_key

            if source_key is None or cv_key is None:
                raise RuntimeError("Workflow did not create expected S3 references.")

            s3.head_object(
                Bucket=aws_settings.documents_bucket_name,
                Key=source_key,
            )
            print("Source object exists: True")

            cv_object = s3.get_object(
                Bucket=aws_settings.documents_bucket_name,
                Key=cv_key,
            )

            cv_text = cv_object["Body"].read().decode("utf-8")

            print("Canonical CV object exists: True")
            print(
                "Canonical CV contains candidate:",
                "Alex Morgan" in cv_text,
            )

    finally:
        print()
        print("Cleaning up synthetic DEV data...")

        prefix = f"orders/{TEST_ORDER_ID}/"

        response = s3.list_objects_v2(
            Bucket=aws_settings.documents_bucket_name,
            Prefix=prefix,
        )

        for item in response.get("Contents", []):
            s3.delete_object(
                Bucket=aws_settings.documents_bucket_name,
                Key=item["Key"],
            )

        table.delete_item(
            Key={"order_id": TEST_ORDER_ID},
        )

        remaining = table.get_item(
            Key={"order_id": TEST_ORDER_ID},
        ).get("Item")

        print("DynamoDB cleanup complete:", remaining is None)

        remaining_objects = s3.list_objects_v2(
            Bucket=aws_settings.documents_bucket_name,
            Prefix=prefix,
        ).get("KeyCount", 0)

        print("S3 cleanup complete:", remaining_objects == 0)


if __name__ == "__main__":
    main()

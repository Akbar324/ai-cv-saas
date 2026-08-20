"""Live synthetic presigned-upload and processing smoke test."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

import boto3  # type: ignore[import-untyped]
import requests
from docx import Document

from backend.models.aws_settings import load_aws_settings
from backend.models.order import Order
from backend.services.cv_workflow import process_uploaded_cv
from backend.services.repository_factory import (
    create_document_repository,
    create_order_repository,
)
from backend.services.upload_service import create_source_upload_target

TEST_ORDER_ID = "smoke-presigned-order-001"
TEST_CUSTOMER_ID = "smoke-presigned-customer-001"

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def create_synthetic_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Alex Morgan")
    document.add_paragraph("Dubai, UAE")
    document.add_paragraph("Quality Engineer")
    document.add_paragraph("Example Engineering LLC")
    document.add_paragraph("2021 - Present")
    document.add_paragraph("Perform technical inspections of telecom infrastructure.")
    document.add_paragraph("Review installation work against engineering standards.")
    document.add_paragraph("Coordinate with field teams to resolve quality issues.")
    document.add_paragraph("Prepare inspection reports.")
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor of Engineering")
    document.add_paragraph("Skills")
    document.add_paragraph("Quality Inspection")
    document.add_paragraph("Telecommunications")
    document.add_paragraph("Python")
    document.add_paragraph("AWS")
    document.save(str(path))


def upload_presigned_post(
    *,
    url: str,
    fields: dict[str, str],
    path: Path,
) -> None:
    with path.open("rb") as file:
        files: dict[str, Any] = {
            "file": (
                path.name,
                file,
                DOCX_CONTENT_TYPE,
            )
        }

        response = requests.post(
            url,
            data=fields,
            files=files,
            timeout=60,
        )

    if response.status_code not in {200, 201, 204}:
        raise RuntimeError(
            f"Presigned upload failed: {response.status_code} {response.text}"
        )


def main() -> None:
    aws_settings = load_aws_settings()

    session = boto3.Session(
        profile_name=aws_settings.profile,
        region_name=aws_settings.region,
    )

    document_repository = create_document_repository(
        aws_settings,
        session=session,
    )
    order_repository = create_order_repository(
        aws_settings,
        session=session,
    )

    s3 = session.client("s3")
    table = session.resource("dynamodb").Table(aws_settings.orders_table_name)

    now = datetime.now(UTC)

    order = Order(
        order_id=TEST_ORDER_ID,
        customer_id=TEST_CUSTOMER_ID,
        target_job_title="Junior Cloud Engineer",
        target_industry="Cloud Computing",
        created_at=now,
        updated_at=now,
    )

    print("Creating synthetic order...")

    table.put_item(
        Item=order.model_dump(mode="json", exclude={"documents"}),
        ConditionExpression="attribute_not_exists(order_id)",
    )

    try:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "synthetic-candidate.docx"
            create_synthetic_docx(source_path)

            print("Creating presigned upload target...")

            target = create_source_upload_target(
                order=order,
                filename=source_path.name,
                content_type=DOCX_CONTENT_TYPE,
                document_repository=document_repository,
                order_repository=order_repository,
            )

            print("Upload key:", target.key)
            print(
                "Upload expires in:",
                target.expires_in_seconds,
                "seconds",
            )

            print("Uploading synthetic DOCX through presigned POST...")

            upload_presigned_post(
                url=target.url,
                fields=target.fields,
                path=source_path,
            )

            print(
                "Uploaded object exists:",
                document_repository.exists(target.key),
            )

            print("Processing uploaded CV...")

            processed = process_uploaded_cv(
                order=order,
                job_description=(
                    "Seeking a junior cloud engineer with AWS knowledge, "
                    "basic Python skills, troubleshooting ability, and "
                    "experience working in technical environments."
                ),
                additional_customer_information=(
                    "Candidate is transitioning toward cloud engineering. "
                    "Do not invent production cloud experience."
                ),
            )

            print()
            print("Processing completed.")
            print("Order status:", processed.order_status.value)
            print(
                "Processing status:",
                processed.processing_status.value,
            )
            print("AI provider:", processed.ai_provider)
            print("AI model:", processed.ai_model)
            print("CV version:", processed.current_cv_version)

            stored = table.get_item(
                Key={"order_id": TEST_ORDER_ID},
            ).get("Item")

            print("DynamoDB record exists:", stored is not None)

            cv_key = processed.documents.current_cv_s3_key

            if cv_key is None:
                raise RuntimeError("Processed order has no canonical CV key.")

            cv_object = s3.get_object(
                Bucket=aws_settings.documents_bucket_name,
                Key=cv_key,
            )

            cv_text = cv_object["Body"].read().decode("utf-8")

            print("Canonical CV object exists: True")
            print(
                "Canonical CV contains candidate:",
                "Alex Morgan" in cv_text,
            )

    finally:
        print()
        print("Cleaning up synthetic DEV data...")

        prefix = f"orders/{TEST_ORDER_ID}/"

        response = s3.list_objects_v2(
            Bucket=aws_settings.documents_bucket_name,
            Prefix=prefix,
        )

        for item in response.get("Contents", []):
            s3.delete_object(
                Bucket=aws_settings.documents_bucket_name,
                Key=item["Key"],
            )

        table.delete_item(
            Key={"order_id": TEST_ORDER_ID},
        )

        remaining_order = table.get_item(
            Key={"order_id": TEST_ORDER_ID},
        ).get("Item")

        remaining_objects = s3.list_objects_v2(
            Bucket=aws_settings.documents_bucket_name,
            Prefix=prefix,
        ).get("KeyCount", 0)

        print(
            "DynamoDB cleanup complete:",
            remaining_order is None,
        )
        print(
            "S3 cleanup complete:",
            remaining_objects == 0,
        )


if __name__ == "__main__":
    main()

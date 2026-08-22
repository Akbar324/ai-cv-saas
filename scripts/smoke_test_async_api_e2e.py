"""Live end-to-end async API smoke test."""

from __future__ import annotations

import time
from pathlib import Path
from typing import Any

import requests
from docx import Document

API_ENDPOINT = "https://lh5fbppgxc.execute-api.me-central-1.amazonaws.com"

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def create_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Alex Morgan")
    document.add_paragraph("Dubai, UAE")
    document.add_paragraph("Quality Engineer")
    document.add_paragraph("Example Engineering LLC")
    document.add_paragraph("2021 - Present")
    document.add_paragraph("Perform technical inspections of telecom infrastructure.")
    document.add_paragraph("Review installation work against engineering standards.")
    document.add_paragraph("Coordinate with field teams to resolve quality issues.")
    document.add_paragraph("Prepare inspection reports.")
    document.add_paragraph("Bachelor of Engineering")
    document.add_paragraph("AWS")
    document.add_paragraph("Python")
    document.save(str(path))


def require_ok(response: requests.Response) -> dict[str, Any]:
    if not response.ok:
        raise RuntimeError(f"HTTP {response.status_code}: {response.text}")

    payload = response.json()

    if not isinstance(payload, dict):
        raise RuntimeError("Expected JSON object response.")

    return payload


def main() -> None:
    path = Path("build/async-api-smoke.docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    create_docx(path)

    print("1. Creating order...")

    order = require_ok(
        requests.post(
            f"{API_ENDPOINT}/orders",
            json={
                "customer_id": "smoke-async-customer",
                "target_job_title": "Junior Cloud Engineer",
                "target_industry": "Cloud Computing",
            },
            timeout=30,
        )
    )

    order_id = str(order["order_id"])

    print("Order ID:", order_id)

    print("2. Requesting upload target...")

    target = require_ok(
        requests.post(
            f"{API_ENDPOINT}/orders/{order_id}/upload-url",
            json={
                "filename": "candidate.docx",
                "content_type": DOCX_CONTENT_TYPE,
            },
            timeout=30,
        )
    )

    print("Object key:", target["object_key"])

    print("3. Uploading DOCX directly to S3...")

    with path.open("rb") as file:
        upload = requests.post(
            target["upload_url"],
            data=target["fields"],
            files={
                "file": (
                    path.name,
                    file,
                    DOCX_CONTENT_TYPE,
                )
            },
            timeout=60,
        )

    if upload.status_code not in {200, 201, 204}:
        raise RuntimeError(f"S3 upload failed: {upload.status_code} {upload.text}")

    print("Upload status:", upload.status_code)

    print("4. Queueing CV processing...")

    queued_response = requests.post(
        f"{API_ENDPOINT}/orders/{order_id}/process",
        json={
            "job_description": (
                "Seeking a junior cloud engineer with AWS knowledge, "
                "basic Python skills, troubleshooting ability, and "
                "experience working in technical environments."
            ),
            "additional_customer_information": (
                "Candidate is transitioning toward cloud engineering. "
                "Do not invent production cloud experience."
            ),
        },
        timeout=30,
    )

    queued = require_ok(queued_response)

    print("Process HTTP status:", queued_response.status_code)
    print("Queue response:", queued)

    if queued_response.status_code != 202:
        raise RuntimeError("Expected HTTP 202 from async process route.")

    print("5. Polling order until processing finishes...")

    deadline = time.monotonic() + 120

    while time.monotonic() < deadline:
        current = require_ok(
            requests.get(
                f"{API_ENDPOINT}/orders/{order_id}",
                timeout=30,
            )
        )

        processing_status = current["processing_status"]
        order_status = current["order_status"]

        print(
            "   processing_status=",
            processing_status,
            " order_status=",
            order_status,
            sep="",
        )

        if processing_status == "succeeded":
            print()
            print("ASYNC WORKFLOW SUCCEEDED")
            print("Order ID:", order_id)
            print("Order status:", order_status)
            print(
                "Current CV key:",
                current["documents"]["current_cv_s3_key"],
            )
            print(
                "AI provider:",
                current["ai_provider"],
            )
            print(
                "AI model:",
                current["ai_model"],
            )
            return

        if processing_status == "failed":
            raise RuntimeError(f"Worker marked order failed: {current}")

        time.sleep(3)

    raise TimeoutError(f"Order did not finish within 120 seconds: {order_id}")


if __name__ == "__main__":
    main()

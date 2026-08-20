"""Tests for CV processing and persistence orchestration."""

from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from backend.models.ai import (
    AIUsage,
    CVOptimizationRequest,
    CVOptimizationResult,
)
from backend.models.cv import CanonicalCV
from backend.models.order import (
    Order,
    OrderStatus,
    ProcessingStatus,
)
from backend.repositories.document_repository import (
    DocumentRepository,
    StoredDocument,
)
from backend.repositories.order_repository import (
    OrderPage,
    OrderRepository,
)
from backend.services.ai_provider import AIProvider
from backend.services.cv_persistence import persist_processed_cv


class FakeAIProvider(AIProvider):
    """Fake AI provider for persistence orchestration tests."""

    def optimize_cv(
        self,
        request: CVOptimizationRequest,
    ) -> CVOptimizationResult:
        cv = CanonicalCV.model_validate(
            {
                "personal_details": {
                    "full_name": "Synthetic Candidate",
                },
                "target_role": {
                    "job_title": request.target_job_title,
                    "industry": request.target_industry,
                    "job_description": request.job_description,
                },
                "skills": ["AWS", "Python"],
            }
        )

        return CVOptimizationResult(
            cv=cv,
            provider="fake-provider",
            model="fake-model",
            usage=AIUsage(
                input_tokens=100,
                output_tokens=50,
                total_tokens=150,
            ),
        )


class FakeDocumentRepository(DocumentRepository):
    """In-memory document repository for orchestration tests."""

    def __init__(self) -> None:
        self.objects: dict[str, bytes] = {}

    def put_file(
        self,
        *,
        key: str,
        path: Path,
        content_type: str,
    ) -> StoredDocument:
        data = path.read_bytes()
        self.objects[key] = data

        return StoredDocument(
            key=key,
            content_type=content_type,
            size_bytes=len(data),
        )

    def put_text(
        self,
        *,
        key: str,
        content: str,
        content_type: str,
    ) -> StoredDocument:
        data = content.encode("utf-8")
        self.objects[key] = data

        return StoredDocument(
            key=key,
            content_type=content_type,
            size_bytes=len(data),
        )

    def get_text(self, key: str) -> str:
        return self.objects[key].decode("utf-8")

    def delete(self, key: str) -> None:
        self.objects.pop(key, None)

    def exists(self, key: str) -> bool:
        return key in self.objects


class FakeOrderRepository(OrderRepository):
    """In-memory order repository for orchestration tests."""

    def __init__(self) -> None:
        self.updated_order: Order | None = None

    def create(self, order: Order) -> None:
        raise NotImplementedError

    def get(self, order_id: str) -> Order | None:
        raise NotImplementedError

    def update(self, order: Order) -> None:
        self.updated_order = order.model_copy(deep=True)

    def list_by_customer(
        self,
        customer_id: str,
        *,
        limit: int = 50,
        next_token: str | None = None,
    ) -> OrderPage:
        raise NotImplementedError

    def list_by_order_status(
        self,
        status: OrderStatus,
        *,
        limit: int = 50,
        next_token: str | None = None,
    ) -> OrderPage:
        raise NotImplementedError

    def list_by_processing_status(
        self,
        status: ProcessingStatus,
        *,
        limit: int = 50,
        next_token: str | None = None,
    ) -> OrderPage:
        raise NotImplementedError

    def list_recent(
        self,
        *,
        limit: int = 50,
        next_token: str | None = None,
    ) -> OrderPage:
        raise NotImplementedError


def create_docx(path: Path) -> Path:
    """Create a synthetic CV fixture."""

    document = Document()
    document.add_paragraph("Synthetic Candidate")
    document.add_paragraph("Quality Engineer")
    document.add_paragraph("AWS Python")

    document.save(str(path))

    return path


def sample_order() -> Order:
    now = datetime(2026, 8, 20, 12, 0, tzinfo=UTC)

    return Order(
        order_id="order-001",
        customer_id="customer-001",
        target_job_title="Cloud Engineer",
        target_industry="Cloud Computing",
        created_at=now,
        updated_at=now,
    )


def test_persist_processed_cv_stores_source_and_canonical_json(
    tmp_path: Path,
) -> None:
    source_path = create_docx(tmp_path / "candidate.docx")

    document_repository = FakeDocumentRepository()
    order_repository = FakeOrderRepository()

    order = persist_processed_cv(
        order=sample_order(),
        source_path=source_path,
        provider=FakeAIProvider(),
        document_repository=document_repository,
        order_repository=order_repository,
        job_description="Seeking AWS and Python knowledge.",
    )

    assert order.documents.source_s3_key == ("orders/order-001/source/original.docx")
    assert order.documents.current_cv_s3_key == ("orders/order-001/cv/v1.json")

    assert document_repository.exists("orders/order-001/source/original.docx")
    assert document_repository.exists("orders/order-001/cv/v1.json")


def test_persist_processed_cv_updates_order_workflow_state(
    tmp_path: Path,
) -> None:
    source_path = create_docx(tmp_path / "candidate.docx")

    document_repository = FakeDocumentRepository()
    order_repository = FakeOrderRepository()

    order = persist_processed_cv(
        order=sample_order(),
        source_path=source_path,
        provider=FakeAIProvider(),
        document_repository=document_repository,
        order_repository=order_repository,
    )

    assert order.current_cv_version == 1
    assert order.processing_status is ProcessingStatus.SUCCEEDED
    assert order.order_status is OrderStatus.HUMAN_REVIEW
    assert order.ai_provider == "fake-provider"
    assert order.ai_model == "fake-model"

    assert order_repository.updated_order is not None
    assert order_repository.updated_order.order_id == "order-001"


def test_persist_processed_cv_increments_existing_version(
    tmp_path: Path,
) -> None:
    source_path = create_docx(tmp_path / "candidate.docx")

    document_repository = FakeDocumentRepository()
    order_repository = FakeOrderRepository()

    order = sample_order()
    order.current_cv_version = 2

    result = persist_processed_cv(
        order=order,
        source_path=source_path,
        provider=FakeAIProvider(),
        document_repository=document_repository,
        order_repository=order_repository,
    )

    assert result.current_cv_version == 3
    assert result.documents.current_cv_s3_key == ("orders/order-001/cv/v3.json")


def test_canonical_json_is_valid_and_contains_candidate(
    tmp_path: Path,
) -> None:
    source_path = create_docx(tmp_path / "candidate.docx")

    document_repository = FakeDocumentRepository()

    persist_processed_cv(
        order=sample_order(),
        source_path=source_path,
        provider=FakeAIProvider(),
        document_repository=document_repository,
        order_repository=FakeOrderRepository(),
    )

    content = document_repository.get_text("orders/order-001/cv/v1.json")

    cv = CanonicalCV.model_validate_json(content)

    assert cv.personal_details.full_name == "Synthetic Candidate"
    assert cv.target_role.job_title == "Cloud Engineer"


class FailingAIProvider(AIProvider):
    """AI provider that always fails."""

    def optimize_cv(
        self,
        request: CVOptimizationRequest,
    ) -> CVOptimizationResult:
        raise RuntimeError("Synthetic AI failure")


def test_persist_processed_cv_marks_processing_before_ai(
    tmp_path: Path,
) -> None:
    source_path = create_docx(tmp_path / "candidate.docx")
    order_repository = FakeOrderRepository()

    persist_processed_cv(
        order=sample_order(),
        source_path=source_path,
        provider=FakeAIProvider(),
        document_repository=FakeDocumentRepository(),
        order_repository=order_repository,
    )

    assert order_repository.updated_order is not None
    assert (
        order_repository.updated_order.processing_status is ProcessingStatus.SUCCEEDED
    )


def test_persist_processed_cv_marks_failed_when_ai_fails(
    tmp_path: Path,
) -> None:
    import pytest

    source_path = create_docx(tmp_path / "candidate.docx")
    order_repository = FakeOrderRepository()

    order = sample_order()

    with pytest.raises(
        RuntimeError,
        match="Synthetic AI failure",
    ):
        persist_processed_cv(
            order=order,
            source_path=source_path,
            provider=FailingAIProvider(),
            document_repository=FakeDocumentRepository(),
            order_repository=order_repository,
        )

    assert order.processing_status is ProcessingStatus.FAILED
    assert order_repository.updated_order is not None
    assert order_repository.updated_order.processing_status is ProcessingStatus.FAILED


def test_failed_processing_does_not_increment_cv_version(
    tmp_path: Path,
) -> None:
    import pytest

    source_path = create_docx(tmp_path / "candidate.docx")

    order = sample_order()

    with pytest.raises(RuntimeError):
        persist_processed_cv(
            order=order,
            source_path=source_path,
            provider=FailingAIProvider(),
            document_repository=FakeDocumentRepository(),
            order_repository=FakeOrderRepository(),
        )

    assert order.current_cv_version == 0
    assert order.documents.current_cv_s3_key is None

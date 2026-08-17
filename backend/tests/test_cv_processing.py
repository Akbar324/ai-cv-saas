"""Tests for the CV processing application service."""

from pathlib import Path

import pytest
from docx import Document

from backend.models.ai import (
    AIUsage,
    CVOptimizationRequest,
    CVOptimizationResult,
)
from backend.models.cv import CanonicalCV
from backend.services.ai_provider import AIProvider
from backend.services.cv_processing import process_cv


class FakeAIProvider(AIProvider):
    """Fake provider used to test orchestration without real AI calls."""

    def __init__(self) -> None:
        self.last_request: CVOptimizationRequest | None = None

    def optimize_cv(
        self,
        request: CVOptimizationRequest,
    ) -> CVOptimizationResult:
        self.last_request = request

        cv = CanonicalCV.model_validate(
            {
                "personal_details": {
                    "full_name": "Test Candidate",
                },
                "target_role": {
                    "job_title": request.target_job_title,
                    "industry": request.target_industry,
                    "job_description": request.job_description,
                },
                "professional_summary": "Optimized synthetic CV.",
                "skills": ["AWS", "Python"],
            }
        )

        return CVOptimizationResult(
            cv=cv,
            provider="fake",
            model="fake-model",
            usage=AIUsage(
                input_tokens=100,
                output_tokens=50,
                total_tokens=150,
            ),
        )


def create_docx(path: Path) -> Path:
    """Create a small DOCX fixture."""

    document = Document()
    document.add_paragraph("Test Candidate")
    document.add_paragraph("Quality Engineer")
    document.add_paragraph("AWS Python")

    document.save(str(path))

    return path


def test_process_cv_extracts_document_and_calls_provider(
    tmp_path: Path,
) -> None:
    path = create_docx(tmp_path / "candidate.docx")
    provider = FakeAIProvider()

    result = process_cv(
        path=path,
        provider=provider,
        target_job_title="Cloud Engineer",
        target_industry="Cloud Computing",
        job_description="Seeking AWS and Python experience.",
        additional_customer_information=(
            "Candidate is transitioning into cloud engineering."
        ),
    )

    assert result.provider == "fake"
    assert result.model == "fake-model"
    assert result.cv.target_role.job_title == "Cloud Engineer"

    assert provider.last_request is not None
    assert provider.last_request.document.filename == "candidate.docx"
    assert provider.last_request.document.file_type == "docx"
    assert "Test Candidate" in provider.last_request.document.text

    assert provider.last_request.target_industry == "Cloud Computing"
    assert provider.last_request.job_description == "Seeking AWS and Python experience."
    assert (
        provider.last_request.additional_customer_information
        == "Candidate is transitioning into cloud engineering."
    )


def test_process_cv_returns_provider_usage_metadata(
    tmp_path: Path,
) -> None:
    path = create_docx(tmp_path / "candidate.docx")
    provider = FakeAIProvider()

    result = process_cv(
        path=path,
        provider=provider,
        target_job_title="Cloud Engineer",
    )

    assert result.usage.input_tokens == 100
    assert result.usage.output_tokens == 50
    assert result.usage.total_tokens == 150


def test_process_cv_with_config_uses_configured_provider(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = create_docx(tmp_path / "candidate.docx")
    fake_provider = FakeAIProvider()

    monkeypatch.setattr(
        "backend.services.cv_processing.load_ai_settings",
        lambda: object(),
    )
    monkeypatch.setattr(
        "backend.services.cv_processing.create_ai_provider",
        lambda settings: fake_provider,
    )

    from backend.services.cv_processing import process_cv_with_config

    result = process_cv_with_config(
        path=path,
        target_job_title="Cloud Engineer",
    )

    assert result.provider == "fake"
    assert fake_provider.last_request is not None
    assert fake_provider.last_request.target_job_title == "Cloud Engineer"
